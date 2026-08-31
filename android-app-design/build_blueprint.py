from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont, ImageFilter
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parent
ASSETS = ROOT / "assets"
PUBLIC = WORKSPACE / "public"
OUTPUT = ROOT / "SAS-Lucknow-Android-App-Blueprint-v1.docx"

INK = "123847"
INK_2 = "1C4B5A"
CREAM = "FBF7EE"
PAPER = "FFFDF8"
GOLD = "C99034"
GOLD_LIGHT = "F1CF82"
SAGE = "DCE9DF"
MUTED = "5B6F77"
LINE = "E4D1A9"
WHITE = "FFFFFF"


def font(size: int, bold: bool = False, serif: bool = False):
    name = "georgiab.ttf" if serif and bold else "georgia.ttf" if serif else "arialbd.ttf" if bold else "arial.ttf"
    return ImageFont.truetype(str(Path("C:/Windows/Fonts") / name), size)


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text_center(draw, box, value, fnt, fill, spacing=4):
    x1, y1, x2, y2 = box
    lines = []
    for part in value.split("\n"):
        lines.extend(wrap(part, max(10, int((x2 - x1) / (fnt.size * 0.55)))))
    total = sum(draw.textbbox((0, 0), line, font=fnt)[3] for line in lines) + spacing * (len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line in lines:
        bx = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + (x2 - x1 - (bx[2] - bx[0])) / 2, y), line, font=fnt, fill=fill)
        y += (bx[3] - bx[1]) + spacing


def arrow(draw, start, end, fill=GOLD, width=5):
    fill = f"#{fill}" if not fill.startswith("#") else fill
    draw.line([start, end], fill=fill, width=width)
    x, y = end
    sx, sy = start
    if abs(x - sx) > abs(y - sy):
        pts = [(x, y), (x - 16 if x > sx else x + 16, y - 10), (x - 16 if x > sx else x + 16, y + 10)]
    else:
        pts = [(x, y), (x - 10, y - 16 if y > sy else y + 16), (x + 10, y - 16 if y > sy else y + 16)]
    draw.polygon(pts, fill=fill)


def make_app_map():
    canvas = Image.new("RGB", (1800, 1160), f"#{CREAM}")
    d = ImageDraw.Draw(canvas)
    d.text((80, 55), "SAS Lucknow Android App", font=font(48, True, True), fill=f"#{INK}")
    d.text((80, 120), "One app, two clear journeys", font=font(24), fill=f"#{MUTED}")

    rounded(d, (75, 205, 820, 1035), 30, f"#{PAPER}", f"#{LINE}", 3)
    rounded(d, (980, 205, 1725, 1035), 30, f"#{INK}")
    d.text((125, 250), "PUBLIC EXPERIENCE", font=font(22, True), fill=f"#{GOLD}")
    d.text((1030, 250), "MEMBER EXPERIENCE", font=font(22, True), fill=f"#{GOLD_LIGHT}")

    public_nodes = [
        ("Home", "Song of Life"), ("Pushpanjali", "Offer & certificate"),
        ("Explore", "Lives, Darshan, events"), ("Join Community", "Create account"),
        ("Support the Work", "Contribution hand-off"),
    ]
    member_nodes = [
        ("Dashboard", "Personal orientation"), ("Darshan", "Six inward spaces"),
        ("Sankalp", "Shared commitments"), ("Yogdaan", "Private contribution record"),
        ("Parichay", "Member identity"),
    ]
    for i, (title, desc) in enumerate(public_nodes):
        y = 330 + i * 132
        rounded(d, (125, y, 770, y + 92), 20, f"#{CREAM}", f"#{LINE}", 2)
        d.text((160, y + 18), title, font=font(24, True, True), fill=f"#{INK}")
        d.text((160, y + 54), desc, font=font(18), fill=f"#{MUTED}")
        if i < len(public_nodes) - 1:
            arrow(d, (448, y + 94), (448, y + 126), width=4)
    for i, (title, desc) in enumerate(member_nodes):
        y = 330 + i * 132
        rounded(d, (1030, y, 1675, y + 92), 20, f"#{INK_2}", f"#{GOLD}", 2)
        d.text((1065, y + 18), title, font=font(24, True, True), fill=f"#{WHITE}")
        d.text((1065, y + 54), desc, font=font(18), fill=f"#{SAGE}")
        if i < len(member_nodes) - 1:
            arrow(d, (1353, y + 94), (1353, y + 126), fill=GOLD_LIGHT, width=4)
    arrow(d, (820, 620), (980, 620), width=7)
    d.text((843, 580), "LOGIN", font=font(18, True), fill=f"#{GOLD}")
    path = ASSETS / "app-map.png"
    canvas.save(path, quality=95)
    return path


def make_journey():
    canvas = Image.new("RGB", (1800, 920), f"#{PAPER}")
    d = ImageDraw.Draw(canvas)
    d.text((80, 45), "The first-day member journey", font=font(46, True, True), fill=f"#{INK}")
    d.text((80, 108), "A simple path from discovery to belonging", font=font(23), fill=f"#{MUTED}")
    steps = [
        ("1", "Discover", "See the vision, videos and current invitation."),
        ("2", "Join", "Create an account with Parichay and password."),
        ("3", "Arrive", "Land on a calm personal dashboard."),
        ("4", "Choose", "Enter Darshan, Sangha, Sankalp or study."),
        ("5", "Return", "Resume saved videos, books and reflections."),
    ]
    x = 70
    for i, (num, title, desc) in enumerate(steps):
        y = 235 + (i % 2) * 285
        rounded(d, (x, y, x + 285, y + 220), 26, f"#{CREAM if i % 2 == 0 else SAGE}", f"#{LINE}", 2)
        d.ellipse((x + 24, y + 24, x + 82, y + 82), fill=f"#{GOLD}")
        text_center(d, (x + 24, y + 24, x + 82, y + 82), num, font(22, True), f"#{WHITE}")
        d.text((x + 24, y + 104), title, font=font(25, True, True), fill=f"#{INK}")
        desc_lines = wrap(desc, 25)
        for j, line in enumerate(desc_lines):
            d.text((x + 24, y + 145 + j * 25), line, font=font(17), fill=f"#{MUTED}")
        if i < len(steps) - 1:
            arrow(d, (x + 288, y + 110), (x + 332, 520 if i % 2 == 0 else 345), width=4)
        x += 340
    d.text((80, 835), "Design principle: one meaningful next step on every screen.", font=font(23, True), fill=f"#{GOLD}")
    path = ASSETS / "member-journey.png"
    canvas.save(path, quality=95)
    return path


def make_architecture():
    canvas = Image.new("RGB", (1800, 1040), f"#{CREAM}")
    d = ImageDraw.Draw(canvas)
    d.text((80, 50), "Recommended technical architecture", font=font(46, True, True), fill=f"#{INK}")
    d.text((80, 112), "Native Android shell, shared SAS Lucknow services", font=font(23), fill=f"#{MUTED}")
    columns = [
        (100, "ANDROID APP", ["Kotlin + Jetpack Compose", "Material 3 components", "Encrypted local preferences", "Media playback + downloads"]),
        (650, "SAS API", ["Existing Node.js services", "JWT/session authentication", "Business rules", "Razorpay verification"]),
        (1200, "CONNECTED SERVICES", ["MongoDB Atlas", "AWS S3", "YouTube playlists", "OpenAI / Savitri Sakhi", "Firebase notifications"]),
    ]
    for x, title, items in columns:
        rounded(d, (x, 235, x + 500, 850), 28, f"#{PAPER if x != 650 else INK}", f"#{GOLD}", 3)
        d.text((x + 35, 280), title, font=font(21, True), fill=f"#{GOLD if x != 650 else GOLD_LIGHT}")
        for i, item in enumerate(items):
            y = 370 + i * 90
            d.ellipse((x + 36, y + 5, x + 58, y + 27), fill=f"#{GOLD_LIGHT}")
            d.text((x + 78, y), item, font=font(20, True if i == 0 else False), fill=f"#{INK if x != 650 else WHITE}")
    arrow(d, (600, 540), (650, 540), width=6)
    arrow(d, (1150, 540), (1200, 540), width=6)
    d.text((615, 505), "HTTPS", font=font(14, True), fill=f"#{GOLD}")
    d.text((1164, 505), "API", font=font(14, True), fill=f"#{GOLD}")
    d.text((100, 930), "No duplicate member database. The app and website remain one ecosystem.", font=font(24, True), fill=f"#{INK}")
    path = ASSETS / "architecture.png"
    canvas.save(path, quality=95)
    return path


def draw_phone(base, x, y, title, eyebrow, variant="light", accent=GOLD, lines=None):
    lines = lines or []
    d = ImageDraw.Draw(base)
    bg = INK if variant == "dark" else PAPER
    fg = WHITE if variant == "dark" else INK
    muted = SAGE if variant == "dark" else MUTED
    rounded(d, (x, y, x + 280, y + 540), 38, "#0B1921", "#09151A", 3)
    rounded(d, (x + 12, y + 16, x + 268, y + 524), 28, f"#{bg}")
    rounded(d, (x + 93, y + 26, x + 187, y + 36), 6, "#101418")
    d.text((x + 30, y + 70), eyebrow.upper(), font=font(10, True), fill=f"#{accent}")
    d.text((x + 30, y + 96), title, font=font(25, True, True), fill=f"#{fg}")
    yy = y + 155
    for kind, text in lines:
        if kind == "hero":
            rounded(d, (x + 28, yy, x + 252, yy + 115), 18, f"#{INK_2 if variant == 'dark' else CREAM}", f"#{GOLD}")
            text_center(d, (x + 42, yy + 12, x + 238, yy + 103), text, font(15, True, True), f"#{WHITE if variant == 'dark' else INK}")
            yy += 132
        elif kind == "button":
            rounded(d, (x + 28, yy, x + 252, yy + 48), 16, f"#{GOLD if variant == 'dark' else INK}")
            text_center(d, (x + 28, yy, x + 252, yy + 48), text, font(13, True), f"#{WHITE}")
            yy += 60
        elif kind == "row":
            rounded(d, (x + 28, yy, x + 252, yy + 62), 14, f"#{INK_2 if variant == 'dark' else WHITE}", f"#{LINE}")
            d.text((x + 42, yy + 20), text, font=font(13, True), fill=f"#{fg}")
            yy += 72
        elif kind == "text":
            for line in wrap(text, 29):
                d.text((x + 30, yy), line, font=font(12), fill=f"#{muted}")
                yy += 19
            yy += 10
    rounded(d, (x + 55, y + 488, x + 225, y + 494), 3, f"#{GOLD_LIGHT if variant == 'dark' else LINE}")


def make_contact_sheet():
    canvas = Image.new("RGB", (1700, 1900), f"#{CREAM}")
    d = ImageDraw.Draw(canvas)
    d.text((70, 45), "First mobile screen family", font=font(46, True, True), fill=f"#{INK}")
    d.text((70, 108), "Representative screens; the interactive board contains the complete set.", font=font(23), fill=f"#{MUTED}")
    screens = [
        ("Home", "THE SONG OF LIFE", "dark", [("hero", "A path towards a more conscious life"), ("button", "Begin exploring"), ("row", "The Song of Savitri")]),
        ("Pushpanjali", "15 AUGUST 2026", "light", [("hero", "Offer a flower to Sri Aurobindo"), ("row", "Choose your pushpa"), ("button", "Offer Pushpanjali")]),
        ("Dashboard", "NAMASTE", "light", [("hero", "Your quiet member space"), ("row", "Continue your journey"), ("row", "My saved moments")]),
        ("Darshan", "SIX INNER SPACES", "dark", [("row", "The Inner Room"), ("row", "Inner Sound"), ("row", "Watch videos"), ("row", "Savitri Sakhi")]),
        ("Inner Room", "SILENCE", "dark", [("hero", "Golden breathing orb"), ("row", "5  ·  10  ·  15  ·  30 min"), ("button", "Begin")]),
        ("Inner Sound", "LISTENING INWARD", "light", [("hero", "Morning · Peace · 10 min"), ("row", "Quiet Aspiration"), ("row", "Mother's Music"), ("button", "Begin listening")]),
        ("Watch", "SONG OF SAVITRI", "light", [("hero", "The Symbol Dawn · Lines 1-5"), ("row", "Watch later · Bookmark"), ("button", "Open video")]),
        ("Sangha", "THE COMMON FIRE", "light", [("hero", "Share a reflection"), ("row", "Latest from the community"), ("row", "Resonates · Reflect · Save")]),
        ("e-Library", "SEARCH & STUDY", "light", [("hero", "Search across books and pages"), ("row", "Sri Aurobindo"), ("row", "The Mother"), ("row", "My bookmarks")]),
        ("Savitri Sakhi", "ENGLISH · हिन्दी", "dark", [("hero", "Ask about a line, book or meaning"), ("row", "Saved conversations"), ("button", "Ask Savitri Sakhi")]),
        ("Sankalp", "SHARED COMMITMENTS", "light", [("hero", "Current work of the Centre"), ("row", "Purpose · Stage · Participants"), ("button", "Offer Seva")]),
        ("Parichay", "MEMBER PROFILE", "light", [("hero", "UC02-2026-000003"), ("row", "Interests & abilities"), ("row", "Communication preferences"), ("button", "Save Parichay")]),
    ]
    for i, (title, eyebrow, variant, lines) in enumerate(screens):
        col = i % 4
        row = i // 4
        draw_phone(canvas, 75 + col * 405, 190 + row * 565, title, eyebrow, variant, lines=lines)
    path = ASSETS / "screen-contact-sheet.png"
    canvas.save(path, quality=95)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=11, color=INK, bold=False, italic=False, serif=False):
    name = "Georgia" if serif else "Calibri"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_body(doc, text, bold_lead=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_run(r, 9, GOLD, True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run(r, 11, INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run(r, {1: 16, 2: 13, 3: 12}[level], INK if level != 2 else GOLD, True, serif=level == 1)
    return p


def add_callout(doc, title, text, fill=CREAM):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title.upper())
    set_run(r, 9, GOLD, True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run(r, 11, INK, True if len(text) < 90 else False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_picture(doc, path, width=6.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = bool(caption)
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    alt_text = caption or Path(path).stem.replace("-", " ").title()
    picture._inline.docPr.set("descr", alt_text)
    picture._inline.docPr.set("title", alt_text)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(caption)
        set_run(r, 9, MUTED, italic=True)


def page_break(doc):
    doc.add_page_break()


def mark_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for level, size, before, after, color in ((1, 16, 18, 10, INK), (2, 13, 12, 6, GOLD), (3, 12, 8, 4, INK_2)):
        style = styles[f"Heading {level}"]
        style.font.name = "Georgia" if level == 1 else "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("SAS LUCKNOW  /  ANDROID APP BLUEPRINT")
    set_run(r, 8, MUTED, True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Sri Aurobindo Society, Lucknow · Product design draft · August 2026")
    set_run(r, 8, MUTED)


def add_screen_card(doc, number, title, purpose, content, primary, linkage):
    add_kicker(doc, f"Screen {number:02d}")
    add_heading(doc, title, 2)
    add_body(doc, purpose, align=WD_ALIGN_PARAGRAPH.LEFT, after=4)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Shows: ")
    set_run(r, 10, GOLD, True)
    r = p.add_run(content)
    set_run(r, 10, INK)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Primary action: ")
    set_run(r, 10, GOLD, True)
    r = p.add_run(primary)
    set_run(r, 10, INK, True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("Leads to: ")
    set_run(r, 10, GOLD, True)
    r = p.add_run(linkage)
    set_run(r, 10, MUTED)


def build_doc(app_map, journey, architecture, contact):
    doc = Document()
    configure_doc(doc)

    # Cover
    add_kicker(doc, "Sri Aurobindo Society · Lucknow Centre")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("SAS Lucknow Android App")
    set_run(r, 30, INK, True, serif=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Product blueprint, user journeys and first screen design system")
    set_run(r, 14, GOLD, False, italic=True, serif=True)
    add_picture(doc, PUBLIC / "song-of-life-banner.png", 6.5)
    add_callout(doc, "The experience promise", "A calm digital home for study, meditation, community and shared work — rooted in The Song of Life.", SAGE)
    add_body(doc, "Prepared for review before Android development begins. This document defines the app structure, page linkages, content behaviour and initial visual direction.", align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    add_kicker(doc, "Executive view")
    add_heading(doc, "What we are building", 1)
    add_body(doc, "The SAS Lucknow Android app will be a native mobile doorway into the same ecosystem already available on saslucknow.in. It will not become a second disconnected platform. A member will use one account, one Parichay, one Yogdaan record and one set of saved content across the website and app.")
    add_callout(doc, "Recommended product position", "Not a smaller website. A quieter, faster and more personal daily companion.")
    add_heading(doc, "Who it serves", 2)
    add_body(doc, "Visitors can discover the Centre, join the community, experience Pushpanjali, watch public videos and support the work. Signed-in members receive a personal dashboard, six Darshan spaces, the Sangha community, Sankalp participation, private Yogdaan history and their Parichay profile.")
    add_heading(doc, "What success feels like", 2)
    for label, text in [
        ("Simple", "A first-time visitor always sees one obvious next step."),
        ("Peaceful", "The visual system feels luminous and composed, never noisy."),
        ("Personal", "Saved study, meditation time, videos and profile return with the member."),
        ("Connected", "Website, Android app and backend share the same source of truth."),
        ("Trustworthy", "Sensitive data and payments are handled transparently and securely."),
    ]:
        add_body(doc, f"{label}. {text}", bold_lead=f"{label}.", align=WD_ALIGN_PARAGRAPH.LEFT, after=4)
    page_break(doc)

    add_kicker(doc, "Information architecture")
    add_heading(doc, "One app, two journeys", 1)
    add_body(doc, "The public journey welcomes and explains. The member journey remembers and deepens. A visitor crosses from public to member space only through a deliberate login or Join the Community action.")
    add_picture(doc, app_map, 6.5, "Figure 1. Public and signed-in sections of the Android app.")
    page_break(doc)

    add_kicker(doc, "Member journey")
    add_heading(doc, "From discovery to belonging", 1)
    add_body(doc, "The app should never present the complete system at once. The home screen introduces the vision; the account flow collects only what is needed; the dashboard then offers a small number of meaningful destinations.")
    add_picture(doc, journey, 6.5, "Figure 2. The recommended first-day experience.")
    add_callout(doc, "Rule for every page", "Show where the member is, what matters now, and one clear next action.", CREAM)
    page_break(doc)

    add_kicker(doc, "Design language")
    add_heading(doc, "The Song of Life, translated to mobile", 1)
    add_body(doc, "The banner's golden sunrise remains the emotional starting point. Inside the app, the palette becomes quieter: deep teal for attention, warm cream for reading, luminous gold for guidance, and soft sage for restorative spaces. Large serif headings create dignity; concise sans-serif controls preserve clarity.")
    add_picture(doc, contact, 6.3, "Figure 3. Representative screens from the first mobile design family.")
    page_break(doc)

    add_kicker(doc, "Navigation")
    add_heading(doc, "A clean mobile hierarchy", 1)
    add_body(doc, "Public screens use a compact top bar and contextual actions. After login, a four-item bottom navigation keeps the most frequent destinations visible: Home, Darshan, Sangha and Profile. Sankalp, Yogdaan and notifications are reached from the dashboard and profile menu. This prevents Darshan's six experiences from competing with member administration.")
    add_heading(doc, "Darshan contains six destinations", 2)
    for title, desc in [
        ("The Inner Room", "Guided silence with a breathing aureole, timer and private reflection."),
        ("Inner Sound", "Curated meditation audio organised by time, state and duration."),
        ("Watch Videos", "Song of Savitri and Gatherings, with bookmarks and saved moments."),
        ("Sangha", "A gentle member feed for reflection, photo, video, artwork and polls."),
        ("e-Library", "Search, read and bookmark pages from the shared study collection."),
        ("Savitri Sakhi", "Bilingual study companion with saved conversations and references."),
    ]:
        add_body(doc, f"{title}. {desc}", bold_lead=f"{title}.", align=WD_ALIGN_PARAGRAPH.LEFT, after=4)
    add_callout(doc, "Navigation safeguard", "Do not duplicate these six destinations as chips, cards and menu items on the same screen. One hierarchy is enough.")
    page_break(doc)

    add_kicker(doc, "Screen designs")
    add_heading(doc, "Public and account screens", 1)
    screens_public = [
        (1, "Splash and secure restore", "Gives the app an immediate identity while securely restoring an existing session.", "Transparent Society logo, subtle golden aura, app name and brief loading state.", "Continue automatically", "Home or Member Dashboard"),
        (2, "Public Home — The Song of Life", "Introduces the vision without reproducing the full website homepage.", "Hero banner, current invitation, Song of Savitri, next Sunday meeting and quick Explore actions.", "Begin exploring", "Explore or Join the Community"),
        (3, "Explore", "Groups public content into understandable paths.", "Sri Aurobindo, The Mother, Darshan Divas, e-Library, events, Sultanpur Shrine and location.", "Open a path", "Selected public detail page"),
        (4, "Pushpanjali offering", "Creates the seasonal flower-offering experience in a mobile-first flow.", "Portrait, offering count, name and email, three flower choices with meanings, selected animation.", "Offer Pushpanjali", "Certificate success"),
        (5, "Pushpanjali certificate", "Makes the certificate instantly useful without overwhelming the user.", "Full certificate preview, download, share, Join the Community and Contribute.", "Join the Community", "Create Member Account, prefilled where consent allows"),
        (6, "Join the Community", "Creates the member account immediately and records Parichay.", "Name, email, mobile, city, interests, password, consent and clear privacy copy.", "Create my member account", "Member Dashboard"),
        (7, "Member Login", "Offers a fast, accessible return path for existing members.", "Email or mobile, password, show-password control and password recovery.", "Enter member space", "Member Dashboard"),
        (8, "Support the Work", "Collects contribution details before a secure payment hand-off.", "Purpose, name, email, mobile, amount, consent, terms and provider status.", "Continue securely", "Razorpay checkout; then receipt and Yogdaan linkage"),
    ]
    for item in screens_public:
        add_screen_card(doc, *item)
    page_break(doc)

    add_kicker(doc, "Screen designs")
    add_heading(doc, "Member core screens", 1)
    screens_member_a = [
        (9, "Member Dashboard", "Welcomes the member and orients them without turning into a metric wall.", "Resume card, next meaningful action, unread notifications, current Sankalp and weekly silence.", "Continue", "The member's last meaningful activity"),
        (10, "Darshan Hub", "Presents the six contemplative spaces as one coherent family.", "Six destination cards, each with a short purpose line and restrained icon.", "Choose a space", "Inner Room, Inner Sound, Watch, Sangha, e-Library or Savitri Sakhi"),
        (11, "Inner Room — setup", "Prepares silence without distraction.", "Matrimandir focus image, 5/10/15/30/open-ended, music choice and volume.", "Begin", "Full-screen meditation session"),
        (12, "Inner Room — session", "Removes normal navigation for the selected period.", "Breathing golden aureole, timer, pause/end and subtle sound control.", "End", "Completion and optional private reflection"),
        (13, "Inner Sound", "Makes contemplative listening feel curated rather than a grid of identical buttons.", "Now playing, Morning/Evening/States/Duration collections, distinct tracks, personal playlists and infinite meditation.", "Begin listening", "Audio player; optionally enter Inner Room"),
        (14, "Watch Library", "Brings Song of Savitri and Gatherings into one searchable member experience.", "Collection tabs, search, My Videos, Watch Later, Bookmarked and thumbnail cards from YouTube.", "Open video", "Video detail"),
        (15, "Video Detail", "Turns viewing into attentive study.", "Player, Savitri reference, transcript search, bookmark, watch later, share, save moment and private note.", "Save this moment", "Timestamped bookmark in My Videos"),
    ]
    for item in screens_member_a:
        add_screen_card(doc, *item)
    page_break(doc)

    add_kicker(doc, "Screen designs")
    add_heading(doc, "Member study, community and service", 1)
    screens_member_b = [
        (16, "Sangha Feed", "Creates a warm, member-only social rhythm without copying the noise of Facebook.", "Latest first, IST timestamps, Reflection/Video/Artwork/Photo/Poll, My Posts, Photos and Videos.", "Share with Sangha", "Composer or post detail"),
        (17, "Create Post or Poll", "Makes posting clear and compact.", "1,000-word limit, optional photo/video for every type; polls allow two to four options with ten words each.", "Publish", "Sangha Feed"),
        (18, "e-Library", "Supports both discovery and return-to-study.", "Full-text search, source filters, collections, recent reading and My Bookmarks.", "Open reading", "Reader with bookmark position"),
        (19, "Savitri Sakhi", "Offers focused bilingual study with citations and saved history.", "English/Hindi question box, book/line lookup, recent conversations, copy/save response and feedback.", "Ask", "Answer with reference and optional source opening"),
        (20, "Sankalp", "Explains collective work before asking for participation or contribution.", "Purpose, stage, participant count, progress narrative, rules and Offer Seva.", "Offer Seva", "Contribution form or participation confirmation"),
        (21, "My Yogdaan", "Keeps contribution amounts private and verifiable.", "Verified contributions, dates, purposes, provider references and downloadable acknowledgements.", "View acknowledgement", "Receipt detail"),
        (22, "Parichay", "Makes member identity useful and editable.", "Unique member ID, contact details, interests, skills, seva preference and notification choices.", "Save Parichay", "Updated member profile"),
        (23, "Notifications", "Surfaces only relevant updates.", "New gallery/video, Sankalp progress, community responses, event reminders and account notices.", "Open update", "The exact originating content"),
    ]
    for item in screens_member_b:
        add_screen_card(doc, *item)
    page_break(doc)

    add_kicker(doc, "Content and behaviour")
    add_heading(doc, "What is shared, saved and private", 1)
    add_heading(doc, "Shared with all visitors", 2)
    add_body(doc, "Public vision content, event information, location, public videos, e-Library catalogue previews, Pushpanjali campaign and public contribution invitation.")
    add_heading(doc, "Shared with signed-in members", 2)
    add_body(doc, "Sangha posts and comments, community notifications, Sankalp participation counts and approved member-only media.")
    add_heading(doc, "Private to the member", 2)
    add_body(doc, "Parichay details, saved study positions, watch history, bookmarks, meditation minutes, private reflections, personal playlists, Savitri Sakhi history and Yogdaan amounts.")
    add_heading(doc, "Administrative only", 2)
    add_body(doc, "Moderation tools, member-role changes, payment verification, content publishing, reports and audit history. The Android member app should not expose the full administrator console in the first release.")
    add_callout(doc, "Privacy promise", "Never turn meditation time, contributions or inner reflections into competitive public metrics.", SAGE)
    page_break(doc)

    add_kicker(doc, "Technology")
    add_heading(doc, "Recommended implementation", 1)
    add_picture(doc, architecture, 6.5, "Figure 4. Android app connected to the existing SAS Lucknow backend.")
    add_heading(doc, "Android foundation", 2)
    add_body(doc, "Use Kotlin and Jetpack Compose with Material 3, a single-activity architecture, ViewModels, coroutines, Retrofit/OkHttp, Room for structured offline data and Media3 for audio/video playback. Use Android Credential Manager for secure sign-in and encrypted storage for tokens.")
    add_heading(doc, "Backend reuse", 2)
    add_body(doc, "Extend the existing Node.js APIs rather than placing business rules in the app. MongoDB Atlas remains the primary data store, S3 stores approved media, YouTube remains the video source, Razorpay verifies payments server-side and Firebase Cloud Messaging delivers notifications.")
    page_break(doc)

    add_kicker(doc, "Connection contract")
    add_heading(doc, "Same ecosystem, mobile-safe access", 1)
    add_body(doc, "The Android app will use saslucknow.in as its single API host and will not connect directly to MongoDB Atlas, S3, Razorpay or OpenAI. Existing data, business rules and provider accounts remain unchanged.")
    add_heading(doc, "Existing services to reuse", 2)
    connection_rows = [
        ("Public content", "/api/savitri-videos, /api/gallery-items, /api/library-search", "Song of Savitri, Gatherings and e-Library"),
        ("Study companion", "/api/savitri-sakhi", "Bilingual Savitri answers through the server"),
        ("Member space", "/api/participation/member/*", "Login, dashboard, Sangha, profile and receipts"),
        ("Participation", "/api/participation/*", "Parichay, Sankalp, public overview and account linkage"),
        ("Payments", "/api/participation/payments/razorpay/*", "Order creation and server-side verification"),
        ("Campaign", "/api/pushpanjali-offerings", "Offering, counter and e-Certificate workflow"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1800, 3900, 3660])
    mark_table_header(table.rows[0])
    for idx, text in enumerate(("Capability", "Current service", "Android use")):
        set_cell_shading(table.rows[0].cells[idx], INK)
        p = table.rows[0].cells[idx].paragraphs[0]
        r = p.add_run(text)
        set_run(r, 9.5, WHITE, True)
    for capability, service, android_use in connection_rows:
        cells = table.add_row().cells
        for i, value in enumerate((capability, service, android_use)):
            set_cell_margins(cells[i])
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[i], CREAM)
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            set_run(r, 8.7, INK, bold=i == 0)
    set_table_geometry(table, [1800, 3900, 3660])
    add_heading(doc, "Required mobile compatibility layer", 2)
    add_body(doc, "The current website protects write requests with same-origin checks and browser cookies. A native Android app has no browser origin, so the backend must add a versioned mobile authentication adapter - for example /api/mobile/v1 - with short-lived bearer access tokens, rotating refresh tokens, device revocation and the same existing member records. This is an access-layer change only; it does not create a second database or duplicate business logic.")
    add_callout(doc, "Connection rule", "Reuse the same backend and databases. Do not embed website cookies, provider secrets or database credentials in the APK.", SAGE)
    page_break(doc)

    add_kicker(doc, "Security and quality")
    add_heading(doc, "Non-negotiable safeguards", 1)
    safeguards = [
        "OpenAI, Razorpay secret and database credentials must never be bundled inside the Android APK.",
        "All payment success states must be confirmed by the SAS server and provider signature verification.",
        "Member tokens must use short-lived access with safe refresh and remote revocation.",
        "Sangha uploads must enforce file type, size, malware checks, S3 access policy and moderation controls.",
        "Private screens must avoid sensitive previews in Android's recent-apps view where appropriate.",
        "Every key flow must work with large text, TalkBack, high contrast and one-handed navigation.",
        "English ships first with Hindi-ready strings; no page should hard-code untranslatable text.",
    ]
    for s in safeguards:
        add_bullet(doc, s)
    add_heading(doc, "Offline behaviour", 2)
    add_body(doc, "Cache the member dashboard, library catalogue, bookmarks and downloaded audio only with explicit consent. Queue drafts safely, clearly label offline states and never imply that a payment or Sangha post succeeded until the server confirms it.")
    page_break(doc)

    add_kicker(doc, "Delivery plan")
    add_heading(doc, "Build in four reviewable phases", 1)
    rows = [
        ("1 · Foundation", "Design system, navigation, authentication, API contract, analytics and accessibility baseline.", "Clickable prototype + technical skeleton"),
        ("2 · Public & identity", "Home, Explore, Pushpanjali, Join Community, Login, support hand-off and notifications permission.", "Internal Android alpha"),
        ("3 · Darshan & content", "Dashboard, Inner Room, Inner Sound, Watch, e-Library and Savitri Sakhi.", "Closed member beta"),
        ("4 · Community & service", "Sangha, Sankalp, Yogdaan, Parichay, hardening, Play Store preparation and monitoring.", "Production release"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1900, 4960, 2500])
    mark_table_header(table.rows[0])
    for idx, text in enumerate(("Phase", "Scope", "Review gate")):
        set_cell_shading(table.rows[0].cells[idx], INK)
        p = table.rows[0].cells[idx].paragraphs[0]
        r = p.add_run(text)
        set_run(r, 10, WHITE, True)
    for phase, scope, gate in rows:
        cells = table.add_row().cells
        for i, value in enumerate((phase, scope, gate)):
            set_cell_margins(cells[i])
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[i], CREAM)
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            set_run(r, 9.5, INK, bold=i == 0)
    set_table_geometry(table, [1900, 4960, 2500])
    add_heading(doc, "First approval checkpoint", 2)
    add_body(doc, "Review the screen family for navigation, tone and content hierarchy before coding begins. Once approved, the next artifact should be a working Jetpack Compose navigation prototype using mock data, followed by API integration screen by screen.")
    add_callout(doc, "Recommended first development sprint", "Splash, Home, Join Community, Login, Member Dashboard and Darshan Hub — with the final design tokens established once.")
    page_break(doc)

    add_kicker(doc, "Definition of done")
    add_heading(doc, "Release acceptance checklist", 1)
    checks = [
        "The same credentials work on website and Android app.",
        "Every navigation destination has a clear back path and meaningful empty state.",
        "No secret is present in the APK or public logs.",
        "Inner Room runs without notifications and restores safely after interruption.",
        "Watch Later, bookmarks, library positions and private reflections sync reliably.",
        "Sangha preserves line breaks, IST timestamps and media behaviour.",
        "Verified contributions appear in My Yogdaan only after server confirmation.",
        "The app passes small-phone, large-text, TalkBack and unstable-network testing.",
        "Crash reporting, privacy disclosures, consent and account deletion paths are ready.",
        "Play Store listing, screenshots, data-safety form and support contact are complete.",
    ]
    for idx, item in enumerate(checks, 1):
        add_body(doc, f"{idx:02d}  {item}", bold_lead=f"{idx:02d}", align=WD_ALIGN_PARAGRAPH.LEFT, after=5)
    add_callout(doc, "Decision requested", "Approve the screen hierarchy and first design direction. Development can then begin with the shared Android design system and authentication shell.", SAGE)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    ASSETS.mkdir(parents=True, exist_ok=True)
    app_map = make_app_map()
    journey = make_journey()
    architecture = make_architecture()
    contact = make_contact_sheet()
    print(build_doc(app_map, journey, architecture, contact))
