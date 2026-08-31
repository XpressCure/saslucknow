from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "android-app-design" / "SAS-Lucknow-Android-App-Web-Parity-Specification-v2.docx"
LOGO = ROOT / "public" / "society-logo-transparent.png"

TEAL = RGBColor(16, 58, 70)
TEAL_2 = RGBColor(29, 81, 92)
GOLD = RGBColor(201, 139, 42)
GOLD_LIGHT = "F7EEDB"
INK = RGBColor(18, 45, 55)
MUTED = RGBColor(91, 111, 117)
WHITE = RGBColor(255, 255, 255)
PALE = "F4F7F6"
LINE = "D8C7A5"


def set_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold, color=color)


def add_table(doc, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[int]):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    for idx, header in enumerate(headers):
        shade(table.rows[0].cells[idx], "103A46")
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE, size=9)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_idx % 2:
                shade(cells[idx], PALE)
            set_cell_text(cells[idx], value)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 if level == 0 else 0.62)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)
    return p


def add_para(doc, text="", *, bold=False, italic=False, color=INK, size=10.5, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def callout(doc, label, text, fill=GOLD_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    # Mark the single callout row so assistive technology announces it as a
    # labelled block rather than an unheaded data table.
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper())
    set_font(r, size=8.5, bold=True, color=GOLD)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.22
    r2 = p2.add_run(text)
    set_font(r2, size=10.4, bold=True, color=INK)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, TEAL, 18, 10),
        ("Heading 2", 13, TEAL_2, 14, 7),
        ("Heading 3", 12, TEAL_2, 10, 5),
    ):
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("SAS LUCKNOW  /  ANDROID WEB-PARITY SPECIFICATION")
    set_font(hr, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Sri Aurobindo Society, Lucknow  |  v2.0  |  Page ")
    set_font(fr, size=8, color=MUTED)
    add_page_field(fp)


def cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(0.85))
        doc_pr = run._r.xpath(".//wp:docPr")
        if doc_pr:
            doc_pr[0].set("descr", "Sri Aurobindo Society transparent symbol")
    add_para(doc, "SRI AUROBINDO SOCIETY · LUCKNOW", bold=True, color=GOLD, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_para(doc, "Android App", bold=True, color=TEAL, size=29, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "Web Parity & Screen Requirements", bold=True, color=TEAL, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Exact-content specification before the screen redesign", italic=True, color=MUTED, size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    callout(doc, "Parity decision", "The Android app will contain the same user-facing fields, copy, actions, states, counters, validations and content as the current website. Only the layout and navigation may adapt for Android screen sizes.")
    add_para(doc, "Prepared from the live website and the current application source", color=MUTED, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "Version 2.0 · 13 August 2026", color=MUTED, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_page_break(doc)


def build():
    doc = Document()
    configure(doc)
    cover(doc)

    add_heading(doc, "Document purpose and approval rule", 1)
    add_para(doc, "This document replaces the earlier simplified app-screen concept. It is the content and behaviour contract for the Android app. No screen design should be finalised until every item in this specification has been traced to a web element and approved.")
    add_table(doc, ["Item", "Decision"], [
        ("Product goal", "Create an Android app for the existing SAS Lucknow public website and member portal."),
        ("Parity standard", "Exact feature, content and state parity with the live web application."),
        ("Permitted Android change", "Responsive rearrangement, native navigation, touch targets, system sharing, notification delivery and device-safe media handling."),
        ("Not permitted", "Omitting, combining or renaming a field, message, counter, action, validation, state or content block without written approval."),
        ("Connections", "Use the same Node.js APIs, MongoDB Atlas data, AWS/S3 media, YouTube playlists, Razorpay workflow and Savitri Sakhi service."),
        ("Current deliverable", "Requirements document only. Exact mobile screens are the next deliverable after approval."),
    ], [2300, 7060])

    add_heading(doc, "1. Product architecture", 1)
    callout(doc, "One platform", "The website and Android app are two clients of the same service layer and data. The app must not create a second database or duplicate content workflow.")
    add_table(doc, ["Layer", "Existing connection retained", "Android requirement"], [
        ("Identity", "Member login and participation/member APIs", "Token-capable mobile session adapter; never store passwords or secrets in source code."),
        ("Data", "MongoDB Atlas", "All member, Parichay, Sankalp, Yogdaan, Sangha and Pushpanjali records remain server-managed."),
        ("Media", "AWS S3 plus YouTube", "Images/media use signed or public server URLs; video playlists reuse the same YouTube IDs and metadata."),
        ("Payments", "Razorpay order and verification APIs", "Use Android checkout only when the contribution facility is enabled; credit Yogdaan only after server verification."),
        ("AI", "Savitri Sakhi API and local Savitri index", "Same English/Hindi answers, line lookup, citations and stored member history."),
        ("Email", "Server-side certificate delivery", "The app shows delivery status but never contains Gmail credentials."),
    ], [1550, 3250, 4560])

    add_heading(doc, "2. Global parity rules", 1)
    for text in [
        "Branding: transparent Society logo, Sri Aurobindo Society name, Lucknow/Gomti Nagar Centre identity, deep teal and golden-aura palette.",
        "Language: retain English/Hindi controls and all bilingual copy already shown on the website.",
        "Audio: retain the meditation music control. Background music pauses whenever a video, media item, modal or Inner Room/Inner Sound session plays.",
        "Navigation: public sections remain discoverable; logged-in navigation keeps Darshan as the parent of exactly six spaces.",
        "States: loading, empty, submitting, success, validation, error, offline and permission-denied states must be designed—not left as generic system messages.",
        "Accessibility: 48dp touch targets, scalable text, meaningful TalkBack labels, keyboard support where relevant, sufficient contrast and reduced-motion support.",
        "Content fidelity: spelling, dates, titles, references, contact details and certificate numbering must come from the same data/configuration as web.",
    ]:
        bullet(doc, text)

    add_page_break(doc)
    add_heading(doc, "3. Public app screen inventory", 1)
    add_para(doc, "The public Android experience mirrors the complete public website, not a shortened promotional landing page.")
    add_table(doc, ["Android screen / section", "Exact web elements to retain"], [
        ("App entry / Home", "Society logo and centre identity; English/Hindi switch; meditation Play/Stop; The Song of Life banner; instant Pushpanjali campaign opening while active."),
        ("The Vision", "Sri Aurobindo Society · Lucknow Centre; THE VISION; four numbered pillars and their explanation: Inner growth, Conscious living, Human unity, Spiritual evolution."),
        ("Lives & Vision", "Sri Aurobindo portrait, 1872–1950, biography and life-sketch link; The Mother/Mirra Alfassa portrait, 1878–1973, biography and life link; rights/source note."),
        ("The Song of Savitri", "Section title and explanatory copy; horizontal video rail; previous/next controls; YouTube playback; Part, Book No., Canto No., Name of Canto, Line Nos., Page No.; full playlist action; add-video action."),
        ("A Living Movement", "Puducherry, Auroville, Sultanpur and Lucknow cards, descriptions and their existing links/actions in the same order."),
        ("e-Library", "Title and explanation; languages; search; All/Books/Audio/Explore filters; all eight collections; result title/category/source/snippet; loading, empty and error states; source attribution."),
        ("Lecture Archive", "Three archive themes, Facebook video module and official Facebook-page action."),
        ("Gatherings through the years", "Section copy; media rail; previous/next; video/photo display; category, event title, description and date; YouTube playlist; add-event-media action."),
        ("Upcoming gatherings", "EVERY SUNDAY; Weekly collective meeting; Quiet reflection, study and fellowship; 6:00–7:00 PM; Plan your visit."),
        ("Visit the Centre", "4/668 Vijayant Khand address; Sunday meeting time; Mr. Rajendra Kumar Singh’s roles; telephone; email; Get directions; Call the centre; embedded map."),
        ("Join the Community", "Create your member account and Member Login choices, with the same supporting copy and destinations."),
        ("Support the work", "Section title and explanation; Contribute thoughtfully action; current ‘facility opening soon’ message whenever payments are intentionally disabled."),
        ("Footer and Savitri Sakhi", "Logo/centre identity, Towards a Life Divine, Wisdom/Events/Email links and legal line; floating Savitri Sakhi launcher with English/Hindi invitation."),
    ], [2500, 6860])

    add_heading(doc, "4. Pushpanjali — exact web-to-app journey", 1)
    callout(doc, "Critical parity screen", "Pushpanjali is not one simplified form. It is a complete multi-state campaign: opening, flower selection, submission, animated offering, confirmation, certificate preview, email delivery and social/community actions.")

    add_heading(doc, "4.1 Campaign entry and modal shell", 2)
    for text in [
        "Open instantly when the campaign is active; closing reveals the persistent Pushpanjali launcher with 15 August 2026.",
        "Close action; top-right live counter formatted for India: ‘N Devotees Offered Pushpanjali’.",
        "Date line: ‘15 AUGUST 2026 - SRI AUROBINDO'S BIRTHDAY DARSHAN’.",
        "Heading: ‘Pushpanjali to Sri Aurobindo’. Supporting line: ‘Offer a flower in gratitude, aspiration and remembrance.’",
        "Sri Aurobindo portrait with golden aura; accessible portrait description; caption ‘Sri Aurobindo’ and ‘1872-1950’.",
        "Presented-by footer: ‘Sri Aurobindo Society, Lucknow - Gomti Nagar Centre (UC-02)’.",
    ]:
        bullet(doc, text)

    add_heading(doc, "4.2 Offering form — every element", 2)
    add_table(doc, ["Element", "Exact requirement"], [
        ("Your name", "Required; placeholder ‘Enter your full name’; maximum 100 characters; device name autofill."),
        ("Email for your certificate", "Required email; placeholder ‘you@example.com’; maximum 180 characters; email autofill."),
        ("Flower legend", "‘Select your pushpa for Pushpanjali’. Three full-width selectable flower cards; Divine Love selected initially."),
        ("Each flower card", "Original flower photograph, flower name, Mother’s significance/meaning and botanical name/variety. Selected card has a visible animation and accessible selected state."),
        ("Privacy", "‘Your email is used only to deliver this certificate. It is not added to a mailing list.’"),
        ("Submit", "‘Offer Pushpanjali & receive certificate →’; disabled/submitting state reads ‘Preparing your offering...’."),
        ("Validation/error", "Inline required/email errors and server error: ‘Your Pushpanjali could not be recorded. Please try again.’ Preserve entered values."),
        ("Abuse protection", "Keep the hidden honeypot/server protections as implementation behaviour; do not expose the ‘Website’ field visually or to TalkBack."),
    ], [2550, 6810])

    add_heading(doc, "4.3 Flower catalogue", 2)
    add_table(doc, ["Flower", "Botanical name / variety", "Spiritual significance / Mother’s message"], [
        ("Divine Love", "Punica granatum - orange-red, double", "A flower that is said to blossom even in the desert."),
        ("Integral Love for the Divine", "Rosa - white", "Pure, complete, irrevocable, a love that gives itself for ever."),
        ("Power of the Supramental Consciousness", "Hibiscus rosa-sinensis ‘Rukmini’ - deep gold, double", "Organising and active, irresistible in its influence."),
    ], [2300, 3000, 4060])

    add_heading(doc, "4.4 Submission and animated offering", 2)
    for text in [
        "The success view must appear immediately after the user submits; certificate generation and email delivery continue in the background.",
        "The chosen flower uses its transparent cut-out—not a circular photograph—and many flowers fall over the portrait and gather at the bottom.",
        "The live devotee counter updates after a valid recorded offering.",
        "No duplicate offering is created by rotation, app backgrounding, retry or a delayed email response.",
    ]:
        bullet(doc, text)

    add_heading(doc, "4.5 Success view — every element", 2)
    add_table(doc, ["Order", "Required content/action"], [
        ("1", "Success symbol and ‘YOUR PUSHPA HAS BEEN OFFERED’."),
        ("2", "‘With gratitude,’ on its own line, followed by the devotee’s name in italics and a full stop."),
        ("3", "Thank-you copy: ‘Thank you for offering your Pushpanjali to Sri Aurobindo. May this gesture of aspiration remain with you.’"),
        ("4", "Certificate Number label and reference; loading placeholder ‘Being prepared...’."),
        ("5", "Email state: being recorded / being sent in background / sent to email / certificate ready / email failed."),
        ("6", "Full certificate image preview shown by default; ‘Preparing your full certificate preview...’ while rendering."),
        ("7", "Share Certificate on WhatsApp; Share on Facebook; Share on Instagram; Download e-Certificate."),
        ("8", "Join the Community and Contribute; both the same size and visual weight as the current approved website layout."),
        ("9", "Contextual sharing notice/error and ‘Return to the website’."),
    ], [900, 8460])

    add_page_break(doc)
    add_heading(doc, "4.6 Certificate image specification", 2)
    add_table(doc, ["Region", "Exact certificate content"], [
        ("Top identity", "Transparent Society logo beside the centred ‘SRI AUROBINDO SOCIETY, LUCKNOW’; ‘GOMTI NAGAR CENTRE (UC-02)’ below; ornate golden frame."),
        ("Portrait panel", "Sri Aurobindo portrait, centre-cropped within its panel; top-left caption ‘Sri Aurobindo’ and ‘1872-1950’; panel height aligns with the certificate-number baseline."),
        ("Heading", "‘Certificate of Pushpanjali’, centred above the main statement."),
        ("Devotee statement", "‘This certifies that’ followed by the devotee name in italic, underlined styling; ‘has lovingly offered Pushpanjali to Sri Aurobindo on his 154th Birthday.’"),
        ("Flower details", "‘Flower Offered’ prominent; BOTANICAL NAME / VARIETY then value; SPIRITUAL SIGNIFICANCE GIVEN BY THE MOTHER then significance; Mother’s explanatory/message line; selected transparent flower image at right."),
        ("Date", "‘15 AUGUST 2026  |  DARSHAN DIVAS’."),
        ("Reference", "‘CERTIFICATE NUMBER: UC02-000001’ pattern, incremented by the server."),
        ("Rendering", "Same aspect ratio, typography hierarchy and content on screen, download, email and share. Preview scales responsively without changing internal font sizes."),
    ], [2200, 7160])

    add_heading(doc, "4.7 Sharing and email copy", 2)
    callout(doc, "All social platforms", "🙏 With gratitude, I have offered Pushpanjali to Sri Aurobindo on his 154th Birthday.\n\nYou too can offer your Pushpanjali and receive a personalised e-Certificate:\nhttps://www.saslucknow.in/?pushpanjali=1\n\nInitiative of: Sri Aurobindo Society, Lucknow, Gomti Nagar Centre (UC-02)")
    add_table(doc, ["Channel", "Required behaviour"], [
        ("WhatsApp", "Share certificate image plus prepared message through Android Sharesheet; both WhatsApp and WhatsApp Business must appear when installed."),
        ("Facebook", "Open native share flow with the message/link; certificate image included where the OS/app permits, otherwise instruct user to attach the visible image."),
        ("Instagram", "Open native share flow with certificate image where supported; copy the standard message for caption/paste fallback."),
        ("Download", "Save/export the PNG with a person-safe filename; show success or recoverable permission error."),
        ("Email", "Email body starts with ‘YOUR PUSHPA HAS BEEN OFFERED’, includes the approved thank-you text, Regards and centre identity, website link, then the certificate image. Do not add a duplicate flower-detail box outside the certificate."),
    ], [1700, 7660])

    add_heading(doc, "5. Join the Community and member access", 1)
    add_para(doc, "Every ‘Join the Community’ action opens the CREATE YOUR MEMBER ACCOUNT / Parichay form directly. The Pushpanjali success action pre-fills and connects the name, email and certificate number without silently creating an account.")
    add_table(doc, ["Screen", "Exact elements"], [
        ("Join the Community page", "Logo/centre header; Home, Member Login and Administrator links; two-step explanation; public Sankalp overview; CREATE YOUR MEMBER ACCOUNT section."),
        ("Account form", "Pushpanjali-connected notice/reference when present; Full name; Mobile; Email optional; City; Areas of interest; Skills you may offer; Seva you would like to explore; Create your member password; consent; validation; Create My Account."),
        ("Password rule", "At least 10 characters with a letter and a number; account can sign in immediately after successful submission."),
        ("Member Login", "Mobile, password, show/hide password, sign in, account activation/first-password state when applicable, errors and help route."),
        ("Session", "Authenticated header: logo/centre, Namaste first name, menu and Sign out. Disabled/expired member states show a clear support path."),
    ], [2500, 6860])

    add_heading(doc, "6. Member portal — exact screen inventory", 1)
    callout(doc, "Darshan structure", "Darshan contains exactly six subspaces: Inner Room, Inner Sound, Sangha, Watch Videos, e-Library and Savitri Sakhi. They appear under Darshan in the side/drawer menu, not as duplicate page-level navigation strips.")
    add_table(doc, ["Screen", "Required web-parity elements and states"], [
        ("Darshan Home", "Namaste heading and supporting copy; Live Sankalp, My Yogdaan and Acknowledgements metrics; Next meaningful action; View Sankalp; ‘A conscious offering’ quote/content."),
        ("Inner Room — setup", "Dark immersive full-screen environment; glowing/breathing Matrimandir orb; Silence; 5/10/15/30 min and Open-ended; volume slider and percent; Begin; Close."),
        ("Inner Room — active", "No navigation/notifications; orb animation; SILENCE; count-down or count-up timer; music volume; End."),
        ("Inner Room — completion", "‘Remain silent for a moment’; optional thought prompt; weekly minutes; Save thought and Continue; no streaks/competition."),
        ("Inner Sound", "Distinct tracks per selection; Now Playing; visual orb; Begin/Pause; mood buttons Silence/Peace/Aspiration/Nature; duration/countdown; volume; Enter The Inner Room; designed Morning, Evening, States and Duration collections."),
        ("Watch Videos", "My Videos tabs: Watch Later and Bookmarked; Continue Watching; All/Song of Savitri/Gatherings filters; search titles, references and transcripts; result count; same Savitri cards as main site without description; exact metadata; video activity counts; Open, Watch Later, Bookmark, Favourite and Playlist."),
        ("Watch detail", "Player; title/collection/reference; Share; Resonates; Save this moment with actual timestamp and note; transcript and search; personal reflection/note; Save note."),
        ("Sangha composer", "Only Reflection, Video, Artwork, Photo and Poll; post text up to 1,000 words; photo/video optional for every type; media limits; Poll two-to-four options, each up to 10 words; IST notice; Share with Sangha."),
        ("Sangha feed", "Newest first; Community, My Posts, Photos and Videos views; preserve user line breaks; author/role/type/IST time; media; poll voting; Resonates; Comment; Share; Save; comments and 300-word reflection form."),
        ("e-Library", "Member-branded banner; search; same collection catalogue/data; loading/empty/error; bookmarks and personal return points in the Android implementation; source-reading behaviour preserved."),
        ("Savitri Sakhi", "Dedicated page; welcome prompts; English/Hindi input; complete conversation history stored for the member; thinking/error states; exact line references verified against the local Savitri index."),
        ("Sankalp", "Stage and funding/support status, title, summary/purpose, progress, received/still-needed/participants, rules, Offer Seva or ‘online offering coming soon’ state."),
        ("My Yogdaan", "Private verified total, offering count, ledger entries, date/reference/amount, acknowledgement action; only verified provider payments appear."),
        ("Parichay", "Member ID; name, mobile, email, linked Pushpanjali reference; City; Areas of interest; Skills; Seva preference; Save Parichay. ID format UC02-20yy-xxxxxx using the server assignment sequence."),
        ("Payment / acknowledgement", "Secure Yogdaan form and Razorpay transition when enabled; verified-success celebration; View My Yogdaan; acknowledgement with logo, receipt, donor, amount, Sankalp, date, payment reference and Print/Share."),
    ], [2350, 7010])

    add_page_break(doc)
    add_heading(doc, "7. API and data traceability", 1)
    add_table(doc, ["Capability", "Current source/service", "Parity rule"], [
        ("Pushpanjali", "/api/pushpanjali-offerings and certificate-email", "Server allocates count/reference; Android retries idempotently and displays every web status."),
        ("Savitri videos", "/api/savitri-videos + SAS_Lko YouTube playlist", "One data model feeds public rail and member Watch cards."),
        ("Gatherings", "/api/gallery-items + Collective Learnings playlist", "One data model feeds public gallery and member Watch cards."),
        ("e-Library", "/api/library-search + collection catalogue", "Same filters/results and source metadata; Android adds member bookmarks without changing search."),
        ("Savitri Sakhi", "/api/savitri-sakhi + local Savitri corpus", "Same response quality, bilingual input and reference verification; authenticated history stored server-side."),
        ("Community", "/api/participation/member/*", "Member auth protects all Sangha writes, votes, comments and saves."),
        ("Member/Sankalp/Yogdaan", "Participation member and admin services", "Do not expose private amounts publicly; only verified payments update Yogdaan."),
        ("Payments", "/api/participation/payments/razorpay/*", "Feature flag controls visibility; no client-side secret; server verifies signature before success."),
        ("Media", "AWS S3 / approved URLs", "No large video binaries on EC2; enforce file type/size before upload."),
    ], [1850, 3150, 4360])

    add_heading(doc, "8. Android screen register", 1)
    add_para(doc, "The final screen board must cover the following distinct states; combining them into one decorative mock-up is not acceptable.")
    screen_rows = [
        ("P01", "Public Home", "Default, EN/HI, music playing/stopped"),
        ("P02", "Public section navigation", "All website sections and More items"),
        ("P03", "Pushpanjali entry", "Ready form with all three full flower cards"),
        ("P04", "Pushpanjali validation/submitting", "Inline error and Preparing your offering"),
        ("P05", "Pushpanjali animated success", "Falling cut-out flowers and complete thank-you"),
        ("P06", "Certificate preview/actions", "Preview plus all seven actions and notices"),
        ("P07", "Join the Community", "Two-step page plus full account form"),
        ("P08", "Member Login", "Default, show password, error, activation"),
        ("P09", "Song of Savitri", "Populated, loading, empty, add-video"),
        ("P10", "e-Library", "Collections, searching, results, empty, error"),
        ("P11", "Gatherings", "Populated rail, empty playlist, add-event"),
        ("P12", "Centre and support", "Map/contact and contribution unavailable/enabled"),
        ("M01", "Member Darshan Home", "Drawer closed/open, dashboard populated/empty"),
        ("M02", "Inner Room", "Setup, active, resting, reflection prompt"),
        ("M03", "Inner Sound", "Playing/paused, collections, duration/mood"),
        ("M04", "Watch Videos", "All, search, My Videos tabs, empty results"),
        ("M05", "Watch detail", "Player, moment, transcript, note"),
        ("M06", "Sangha", "Composer by type, poll builder, feed, My Posts/media"),
        ("M07", "Member e-Library", "Search, results and saved/bookmarked state"),
        ("M08", "Savitri Sakhi", "Welcome, chat, thinking, referenced answer, history"),
        ("M09", "Sankalp", "Cards, rules, offering unavailable/enabled"),
        ("M10", "My Yogdaan", "Empty, ledger, verified-success, acknowledgement"),
        ("M11", "Parichay", "Assigned/pending member ID and editable profile"),
        ("M12", "Global account states", "Offline, session expired, disabled, sign out"),
    ]
    add_table(doc, ["ID", "Screen", "Required variants"], screen_rows, [700, 2600, 6060])

    add_heading(doc, "9. Responsive adaptation rules", 1)
    for text in [
        "Exact content does not mean shrinking the desktop page. Use single-column stacking, bottom sheets and drawers while preserving all elements.",
        "Pushpanjali portrait remains centred with no blank side band; flower names, meanings and botanical names remain fully readable at 360dp width.",
        "Certificate preview remains a single image; the surrounding controls stack into equal-width rows without changing certificate typography.",
        "Video rails support horizontal swipe plus visible navigation/accessibility controls; media uses 16:9 containers and never overflows.",
        "Member side navigation becomes a drawer; Darshan’s six children stay nested beneath Darshan.",
        "Forms use labelled inputs above controls, correct Android keyboards, autofill hints, focus order and keyboard-safe scrolling.",
        "Test portrait widths 360, 390 and 412dp; landscape/tablet widths 600 and 840dp; font scales 100%, 130% and 200%. No horizontal page scroll.",
    ]:
        bullet(doc, text)

    add_heading(doc, "10. Acceptance checklist", 1)
    checks = [
        "Every live web field, label, helper text, button, link, counter and state is mapped to one Android component.",
        "Pushpanjali ready, submitting, success, certificate, email and share states match this specification exactly.",
        "All three flower cards contain image, full flower name, Mother’s significance/meaning and botanical name/variety.",
        "All public sections and the complete member portal are reachable—none are replaced by a generic ‘Explore’ screen.",
        "Darshan contains exactly six nested spaces and no duplicate page-level navigation strip.",
        "Watch search, Watch Later, Bookmarked, timestamp moments, transcript search and notes work with real website data.",
        "Sangha preserves line breaks, records IST, enforces media/poll/word limits and shows newest posts first.",
        "Parichay shows server-assigned member ID and the data collected during account creation.",
        "Yogdaan changes only after verified payment; feature-disabled messages never open Razorpay.",
        "Same-backend contract is demonstrated in integration tests; no secrets are bundled in the APK.",
        "Visual comparisons pass on 360/390/412dp devices with no clipped text, blank colour strips, overflow or unreadable controls.",
        "TalkBack order and names are correct; reduced motion keeps all content and offering feedback understandable.",
    ]
    for item in checks:
        bullet(doc, "☐ " + item)

    add_heading(doc, "11. Next deliverable after document approval", 1)
    add_para(doc, "After approval, the screen set will be rebuilt from this register. It will show the real website copy and all elements—starting with the complete Pushpanjali ready, submission, success and certificate screens—at Android phone dimensions. Each design will carry its screen ID so it can be checked against this document.")
    callout(doc, "Approval gate", "Do not proceed to UI finalisation or Android implementation if any live website element is absent from this parity specification. Add the missing item here first, then update the screen register.")

    doc.core_properties.title = "SAS Lucknow Android App - Web Parity & Screen Requirements"
    doc.core_properties.subject = "Exact web-to-Android content, screen, state and integration specification"
    doc.core_properties.author = "Sri Aurobindo Society, Lucknow"
    doc.core_properties.keywords = "SAS Lucknow, Android, web parity, Pushpanjali, member portal"
    doc.core_properties.comments = "Prepared for screen redesign approval."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
